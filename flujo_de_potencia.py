import streamlit as st
import pandas as pd
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import google.generativeai as genai

# --- 1. CONFIGURACIÓN DE LA PÁGINA ---
st.set_page_config(page_title="Plataforma Flujo de Potencia", page_icon="⚡", layout="wide")

# --- 2. CONFIGURACIÓN DE LA IA (CON PARCHE ANTBG) ---
model = None
if "GEMINI_API_KEY" in st.secrets:
    try:
        genai.configure(api_key=st.secrets["GEMINI_API_KEY"].strip())
        model = genai.GenerativeModel('gemini-3-flash-preview')
    except Exception as e:
        st.error(f"Error al configurar la IA desde Secrets: {e}")
else:
    # Si la interfaz de Streamlit está mocha y no muestra los Secrets, habilita el plan B manual
    api_key_manual = st.text_input("🔑 Streamlit Cloud no cargó los Secrets. Ingresa tu Gemini API Key manualmente aquí:", type="password")
    if api_key_manual:
        try:
            genai.configure(api_key=api_key_manual.strip())
            model = genai.GenerativeModel('gemini-3-flash-preview')
        except Exception as e:
            st.error(f"Error con la clave ingresada: {e}")
            st.stop()
    else:
        st.warning("⚠️ Ingrese su API Key en el recuadro superior para habilitar el motor de análisis de la IA.")

# --- 3. LÓGICA DE EXTRACCIÓN Y CONSTRUCCIÓN DE TABLAS ---
def procesar_bloque_tabla(df, keyword, doc):
    """
    Busca las columnas que contienen la palabra clave, extrae el bloque de datos
    y construye de forma limpia la tabla nativa dentro de Word de manera secuencial.
    """
    headers = df.columns.astype(str).tolist()
    columnas_coincidentes = [col for col in headers if keyword in col]
    
    if not columnas_coincidentes:
        return None
        
    df_bloque = df[columnas_coincidentes].dropna(how='all').reset_index(drop=True)
    if df_bloque.empty:
        return None

    azul_corporativo = RGBColor(0, 76, 95)
    nombre_seccion = keyword.replace(chr(10), ' ')
    
    # Agregar título de la sub-tabla en el Word
    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run(f"Resultados: {nombre_seccion}")
    run_sub.font.name = 'Ubuntu'
    run_sub.font.size = Pt(14)
    run_sub.font.bold = True
    run_sub.font.color.rgb = azul_corporativo
    
    # Crear tabla nativa en Word
    tabla_word = doc.add_table(rows=len(df_bloque) + 1, cols=len(df_bloque.columns))
    tabla_word.style = 'Table Grid'
    
    # Escribir los encabezados del bloque
    for j, col_name in enumerate(df_bloque.columns):
        cell = tabla_word.cell(0, j)
        cell.text = str(col_name).replace(chr(10), ' ')
        if cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Ubuntu'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = azul_corporativo
            
    # Escribir las filas de datos técnicos
    for i, row in enumerate(df_bloque.itertuples(index=False)):
        for j, val in enumerate(row):
            cell = tabla_word.cell(i + 1, j)
            
            if isinstance(val, float):
                if "p.u." in df_bloque.columns[j] or "Tensión" in df_bloque.columns[j]:
                    cell.text = f"{val:.1f}"  # Redondeo de tensión a 1 decimal solicitado
                else:
                    cell.text = f"{val:.2f}"
            else:
                cell.text = str(val) if pd.notna(val) else ""
                
            if cell.paragraphs[0].runs:
                run_cell = cell.paragraphs[0].runs[0]
                run_cell.font.name = 'Ubuntu'
                run_cell.font.size = Pt(10)
                run_cell.font.color.rgb = azul_corporativo
                
    doc.add_paragraph()
    return df_bloque.to_markdown(index=False)

# --- 4. INTERFAZ DE USUARIO ---
st.title("⚡ Plataforma Flujo de Potencia")
st.write("Potenciado por Gemini 3 Flash Preview")
st.markdown("---")

uploaded_file = st.file_uploader("Selecciona el archivo Excel de Resultados (EFP_RES_...)", type=["xlsx"])

if uploaded_file is not None:
    try:
        excel_file = pd.ExcelFile(uploaded_file)
        sheet_names = excel_file.sheet_names
        st.success(f"📂 Archivo cargado correctamente. Se detectaron {len(sheet_names)} escenarios de estudio.")
    except Exception as e:
        st.error(f"Error al leer el archivo Excel: {e}")
        st.stop()

    if st.button("🚀 Ejecutar Análisis y Generar Informe"):
        
        doc = Document()
        
        # Configurar formato global
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Ubuntu'
        font.size = Pt(18)
        font.color.rgb = RGBColor(0, 76, 95)
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        total_sheets = len(sheet_names)
        contexto_para_conclusion = ""

        # Procesar cada hoja del Excel (Escenarios independientes)
        for index, sheet_name in enumerate(sheet_names):
            status_text.write(f"Procesando Escenario: **{sheet_name}** ({index + 1}/{total_sheets})...")
            
            # Título del Escenario en Word
            p_title = doc.add_paragraph()
            run_title = p_title.add_run(f"Resultados {sheet_name}")
            run_title.font.size = Pt(18)
            run_title.font.bold = True
            run_title.font.name = 'Ubuntu'
            run_title.font.color.rgb = RGBColor(0, 76, 95)
            
            pPr = p_title._p.get_or_add_pPr()
            pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFFFFF"/>'))
            
            df_sheet = pd.read_excel(uploaded_file, sheet_name=sheet_name)
            
            # Procesar e inyectar de forma nativa cada sub-tabla para que no se alteren
            contexto_escenario = ""
            for kw in ["Línea\n[MVA]", "Línea\n[kA]", "Transformador", "Barra"]:
                tabla_md = procesar_bloque_tabla(df_sheet, kw, doc)
                if tabla_md:
                    contexto_escenario += f"\n{tabla_md}\n"
            
            # --- SOLICITUD DE ANÁLISIS A GEMINI (SÓLO SI EL MODELO ESTÁ ACTIVADO) ---
            if model is not None:
                prompt_escenario = f"""
                Actúa como un Ingeniero Consultor Senior en Sistemas Eléctricos de Potencia en Chile.
                Analiza los siguientes resultados del escenario '{sheet_name}'.
                
                {contexto_escenario}
                
                Redacta un análisis técnico ejecutivo de máximo 2 párrafos sobre este escenario.
                Criterios normativos:
                1. Cargabilidades de Líneas y Transformadores deben ser <= 100%.
                2. Voltajes en barras (< 200 kV) en estado estacionario deben estar entre 0.93 y 1.07 p.u.
                """
                try:
                    response = model.generate_content(prompt_escenario)
                    analisis_texto = response.text
                except Exception as e:
                    analisis_texto = f"Error al generar el análisis técnico: {e}"
            else:
                analisis_texto = "Análisis omitido. No se ingresó una API Key válida para ejecutar el motor de IA."

            # Pegar la respuesta del análisis abajo de las tablas físicas del escenario
            p_analisis = doc.add_paragraph()
            run_analisis = p_analisis.add_run(f"Evaluación Técnica:\n{analisis_texto}")
            run_analisis.font.name = 'Ubuntu'
            run_analisis.font.size = Pt(18)
            run_analisis.font.color.rgb = RGBColor(0, 76, 95)
            p_analisis.paragraph_format.space_after = Pt(24)
            
            pPr_a = p_analisis._p.get_or_add_pPr()
            pPr_a.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFFFFF"/>'))
            
            contexto_para_conclusion += f"\n--- Escenario {sheet_name} ---\n{analisis_texto}\n"
            progress_bar.progress((index + 1) / total_sheets)
            
        # --- GENERAR CONCLUSIÓN GENERAL DE CIERRE ---
        if model is not None:
            status_text.write("Redactando Conclusión General del Estudio...")
            prompt_conclusion = f"Basado en los siguientes análisis de escenarios, redacta las Conclusiones y Recomendaciones Generales en un tono altamente corporativo:\n{contexto_para_conclusion}"
            try:
                response_conclusion = model.generate_content(prompt_conclusion)
                conclusion_texto = response_conclusion.text
            except Exception as e:
                conclusion_texto = f"No se pudo compilar la conclusión automáticamente: {e}"
        else:
            conclusion_texto = "Conclusión general no generada por falta de credenciales de IA."
            
        doc.add_page_break()
        p_c_title = doc.add_paragraph()
        run_c_title = p_c_title.add_run("Conclusiones y Recomendaciones Generales")
        run_c_title.font.size = Pt(18)
        run_c_title.font.bold = True
        run_c_title.font.name = 'Ubuntu'
        run_c_title.font.color.rgb = RGBColor(0, 76, 95)
        
        p_c_text = doc.add_paragraph()
        run_c_text = p_c_text.add_run(conclusion_texto)
        run_c_text.font.name = 'Ubuntu'
        run_c_text.font.size = Pt(18)
        run_c_text.font.color.rgb = RGBColor(0, 76, 95)
        
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        status_text.success("✨ ¡Informe técnico de Flujo de Potencia generado con éxito!")
        
        st.download_button(
            label="📥 Descargar Informe Técnico (.docx)",
            data=docx_buffer,
            file_name="Informe_Tecnico_Flujo_de_Potencia.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

st.markdown("---")
st.caption("© 2026 Plataforma Flujo de Potencia - Ingeniería Eléctrica")
